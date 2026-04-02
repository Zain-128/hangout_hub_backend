#!/usr/bin/env python3
"""Extract docsss.pdf text with heading detection; emit DOCX + PDF (28pt headings, 14pt body)."""

import re
import sys
from xml.sax.saxutils import escape as xml_escape

import fitz
from docx import Document
from docx.shared import Mm, Pt
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

# Original PDF uses ~12pt body, >=~16.8pt for headings
HEADING_MIN_SIZE = 14.0

FOOTER_RE = re.compile(r"^\s*--\s*\d+\s+of\s+\d+\s*--\s*$", re.I)


def extract_structured(pdf_path: str):
    doc = fitz.open(pdf_path)
    out: list[tuple[bool, str]] = []
    for pno in range(len(doc)):
        page = doc[pno]
        d = page.get_text("dict")
        blocks = [b for b in d["blocks"] if b.get("type") == 0]
        blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))
        for b in blocks:
            for line in b["lines"]:
                spans = line.get("spans") or []
                if not spans:
                    continue
                text = "".join(s.get("text", "") for s in spans).strip()
                if not text:
                    continue
                if FOOTER_RE.match(text):
                    continue
                max_size = max(float(s.get("size", 12)) for s in spans)
                is_heading = max_size >= HEADING_MIN_SIZE
                out.append((is_heading, text))
    doc.close()
    return out


def write_docx(items: list[tuple[bool, str]], path: str) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Mm(20)
    section.left_margin = section.right_margin = Mm(20)

    for is_heading, text in items:
        para = document.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(28 if is_heading else 14)
        if is_heading:
            run.bold = True
    document.save(path)


def write_pdf(items: list[tuple[bool, str]], path: str) -> None:
    base = getSampleStyleSheet()
    heading_style = ParagraphStyle(
        name="H28",
        parent=base["Normal"],
        fontName="Helvetica-Bold",
        fontSize=28,
        leading=34,
        spaceAfter=10,
    )
    body_style = ParagraphStyle(
        name="B14",
        parent=base["Normal"],
        fontName="Helvetica",
        fontSize=14,
        leading=17,
        spaceAfter=6,
        alignment=TA_JUSTIFY,
    )

    def rl_para(txt: str, style) -> Paragraph:
        safe = xml_escape(txt).replace("\n", "<br/>")
        return Paragraph(safe, style)

    story = []
    for is_heading, text in items:
        story.append(rl_para(text, heading_style if is_heading else body_style))
        story.append(Spacer(1, 2 * mm))

    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    doc.build(story)


def main():
    src = "/Users/tmdev/Downloads/docsss.pdf"
    out_docx = "/Users/tmdev/Downloads/docsss_formatted.docx"
    out_pdf = "/Users/tmdev/Downloads/docsss_formatted.pdf"
    if len(sys.argv) >= 2:
        src = sys.argv[1]
    if len(sys.argv) >= 3:
        out_docx = sys.argv[2]
    if len(sys.argv) >= 4:
        out_pdf = sys.argv[3]

    items = extract_structured(src)
    if not items:
        sys.stderr.write("No text extracted.\n")
        sys.exit(1)
    write_docx(items, out_docx)
    write_pdf(items, out_pdf)
    print(f"Wrote {len(items)} blocks ->\n  {out_docx}\n  {out_pdf}")


if __name__ == "__main__":
    main()
