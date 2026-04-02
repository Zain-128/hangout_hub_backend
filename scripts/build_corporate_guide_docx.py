#!/usr/bin/env python3
"""Build professional Word doc from docsss.pdf: Times New Roman, headings 20pt, body 14pt."""

import re
import sys
from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

HEADING_MIN_SIZE = 14.0
FOOTER_RE = re.compile(r"^\s*--\s*\d+\s+of\s+\d+\s*--\s*$", re.I)


def _set_run_font(run, size_pt: float, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")


def extract_structured(pdf_path: str) -> list[tuple[bool, str]]:
    doc = fitz.open(pdf_path)
    out: list[tuple[bool, str]] = []
    for pno in range(len(doc)):
        page = doc[pno]
        d = page.get_text("dict")
        blocks = [b for b in d["blocks"] if b.get("type") == 0]
        blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))
        for b in blocks:
            for line in b["lines"]:
                spans = line.get("spans") or []
                if not spans:
                    continue
                text = "".join(s.get("text", "") for s in spans).strip()
                if not text:
                    continue
                if FOOTER_RE.match(text):
                    continue
                max_size = max(float(s.get("size", 12)) for s in spans)
                is_heading = max_size >= HEADING_MIN_SIZE
                out.append((is_heading, text))
    doc.close()
    return out


def build_docx(items: list[tuple[bool, str]], path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    first = True
    for is_heading, text in items:
        para = document.add_paragraph()
        if is_heading:
            run = para.add_run(text)
            _set_run_font(run, 20, bold=True)
            para.paragraph_format.space_before = Pt(10 if not first else 0)
            para.paragraph_format.space_after = Pt(8)
            para.paragraph_format.keep_with_next = True
            if first and "corporate" in text.lower() and "portal" in text.lower():
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_after = Pt(18)
        else:
            run = para.add_run(text)
            _set_run_font(run, 14, bold=False)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.15
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        first = False

    document.core_properties.title = "Corporate portal: detailed screen-by-screen guide"
    document.core_properties.subject = "ic-frontend-corporate documentation"
    document.save(str(path))


def main() -> None:
    src = Path("/Users/tmdev/Downloads/docsss.pdf")
    out = Path("/Users/tmdev/Downloads/Corporate_portal_screen_by_screen_guide.docx")
    workspace_copy = Path(
        "/Users/tmdev/Documents/Zain/hangout_hub_backend/docs/Corporate_portal_screen_by_screen_guide.docx"
    )

    if len(sys.argv) >= 2:
        src = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        out = Path(sys.argv[2])

    if not src.is_file():
        print(f"Source PDF not found: {src}", file=sys.stderr)
        sys.exit(1)

    items = extract_structured(str(src))
    if not items:
        print("No text extracted.", file=sys.stderr)
        sys.exit(1)

    out.parent.mkdir(parents=True, exist_ok=True)
    build_docx(items, out)
    print(f"Wrote {len(items)} paragraphs -> {out}")

    workspace_copy.parent.mkdir(parents=True, exist_ok=True)
    build_docx(items, workspace_copy)
    print(f"Copy -> {workspace_copy}")


if __name__ == "__main__":
    main()
